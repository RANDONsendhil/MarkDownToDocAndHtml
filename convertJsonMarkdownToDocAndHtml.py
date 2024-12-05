import json
from docx import Document
import re
import markdown
from bs4 import BeautifulSoup
import os


class ConvertGsscData:
    ret = ""
    json_data = ""

    def __init__(self, **kwargs):
        self.input_file_path = kwargs.get("input_file_path", "")
        self.output_file_path = kwargs.get("output_file_path", "")
        self.appName = kwargs.get("appName", "")

    def get_files_list(self, folder_path):
        return [
            f
            for f in os.listdir(folder_path)
            if os.path.isfile(os.path.join(folder_path, f))
        ]

    def extract_data(self, output_file):
        added_titles = set()
        for item in self.json_data:
            if item["appName"] not in added_titles:
                added_titles.add(item["appName"])
                self.filter_and_sort_by_appname(item["appName"])
        html_content = markdown.markdown(self.get_markdown_content())
        self.convert_to_html_content(html_content, f"{output_file}.html")
        self.convert_to_doc(html_content, f"{output_file}.docx")

    def check_pdf_field(self, filepath):
        try:
            with open(filepath, "r", encoding="utf-8") as file:
                return json.load(file)
        except json.JSONDecodeError as e:
            print(f"Error parsing JSON: {e}")
            return []

    def remove_characters(self, text):
        return re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", text)

    def filter_and_sort_by_appname(self, app_name):
        extracted_data = []
        for entry in self.json_data:
            if entry["appName"] == app_name:
                nom = entry.get("nom", "")
                extracted_data.append(
                    {
                        "appName": entry["appName"],
                        "nom": nom,
                        "contenu": entry["contenu"],
                    }
                )
        sorted_data = sorted(extracted_data, key=lambda x: x["nom"])
        for item in sorted_data:
            nom = item.get("nom", "")
            self.ret += (
                f"\n#{self.remove_characters(item['appName'])}"
                f"\n# {nom}"
                f"\n{self.remove_characters(item['contenu'])}"
            )

    def remove_invalid_xml_chars(self, s):
        return re.sub(r"[\x00-\x08\x0B-\x0C\x0E-\x1F]", "", s)

    def convert_to_doc(self, markdown_content, output_file):
        html_content = markdown.markdown(markdown_content)
        soup = BeautifulSoup(html_content, "html.parser")
        doc = Document()
        for element in soup.children:
            try:
                if element.name == "h1":
                    doc.add_heading(element.get_text(), 0)
                elif element.name == "h2":
                    doc.add_heading(element.get_text(), level=2)
                elif element.name == "h3":
                    doc.add_heading(element.get_text(), level=3)
                elif element.name == "p":
                    doc.add_paragraph(element.get_text())
                elif element.name == "ul":
                    for li in element.find_all("li"):
                        doc.add_paragraph(f"- {li.get_text()}", style="List Bullet")
                elif element.name == "ol":
                    for li in element.find_all("li"):
                        doc.add_paragraph(f"{li.get_text()}", style="List Number")
                elif element.name == "a":
                    text = element.get_text()
                    link = element.get("href")
                    doc.add_paragraph(f"{text} ({link})")
            except ReferenceError:
                print("Value error")
        print(output_file)
        doc.save(output_file)

    def convert_to_html_content(self, html_content, output_html):
        html_page = f"""
      <!DOCTYPE html>
      <html lang="en">
      <head>
          <meta charset="UTF-8">
          <meta http-equiv="X-UA-Compatible" content="IE=edge">
          <meta name="viewport" content="width=device-width, initial-scale=1.0">
          <title>Markdown to HTML Page</title>
      </head>
      <body>
          {html_content}
      </body>
      </html>
      """
        with open(output_html, "w", encoding="utf-8") as f:
            f.write(html_page)
        print(f"Markdown has been converted to an HTML page: {output_html}")

    def run(self):
        try:
            self.json_data = self.check_pdf_field(
                f"{self.input_file_path}/{self.appName}.json"
            )
            self.extract_data(f"{self.output_file_path}/{self.appName.split('.')[0]}")
        except FileNotFoundError:
            print(f"The file '{self.appName}' was not found!")
